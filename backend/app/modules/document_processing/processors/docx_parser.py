"""DOCX parser using python-docx."""
from docx import Document
from typing import List, Optional
import uuid
from app.modules.document_processing.processors.base_parser import BaseParser
from app.modules.document_processing.schemas import (
    CanonicalDocument,
    ParagraphSchema,
    SectionSchema,
    TableSchema,
    DocumentMetadataSchema,
    ProcessingInformation,
    DocumentStatistics
)
from app.modules.document_processing.enums import DocumentLanguage
from app.modules.document_processing.exceptions import ParserException


class DOCXParser(BaseParser):
    """DOCX parser using python-docx."""
    
    def validate_file(self) -> bool:
        """Validate DOCX file."""
        if self.get_file_extension() != '.docx':
            raise ParserException(f"Invalid file type for DOCX parser: {self.get_file_extension()}")
        
        try:
            Document(str(self.file_path))
            return True
        except Exception as e:
            raise ParserException(f"Invalid DOCX file: {str(e)}")
    
    def parse(self) -> CanonicalDocument:
        """Parse DOCX document and return canonical document object.
        
        Returns:
            CanonicalDocument: Structured document object
            
        Raises:
            ParserException: If parsing fails
        """
        try:
            doc = Document(str(self.file_path))
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Extract text content
            paragraphs, sections = self._extract_text_content(doc)
            
            # Extract tables
            tables = self._extract_tables(doc)
            
            # Calculate statistics
            statistics = self._calculate_statistics(paragraphs, tables, sections)
            
            # Build canonical document
            document = CanonicalDocument(
                id=str(uuid.uuid4()),
                document_id=self.document_id,
                title=metadata.title or self.extract_title_from_filename(),
                language=metadata.language,
                page_count=metadata.page_count,
                sections=sections,
                paragraphs=paragraphs,
                tables=tables,
                images=[],  # DOCX images could be extracted but keeping simple for now
                metadata=metadata,
                statistics=statistics,
                raw_text=self._extract_raw_text(doc),
                normalized_text=self._normalize_text(paragraphs),
                processing_information=ProcessingInformation(
                    parser_used=self.__class__.__name__,
                    ocr_used=False
                )
            )
            
            return document
            
        except Exception as e:
            raise ParserException(f"Failed to parse DOCX: {str(e)}")
    
    def _extract_metadata(self, doc: Document) -> DocumentMetadataSchema:
        """Extract DOCX metadata.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            DocumentMetadataSchema: Document metadata
        """
        core_props = doc.core_properties
        
        return DocumentMetadataSchema(
            title=core_props.title,
            author=core_props.author,
            creation_date=core_props.created,
            modification_date=core_props.modified,
            language=DocumentLanguage.UNKNOWN,
            page_count=None,  # DOCX doesn't have pages
            subject=core_props.subject,
            keywords=core_props.keywords.split(',') if core_props.keywords else []
        )
    
    def _extract_text_content(self, doc: Document) -> tuple[List[ParagraphSchema], List[SectionSchema]]:
        """Extract text content with heading detection.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Tuple of (paragraphs, sections)
        """
        paragraphs = []
        sections = []
        paragraph_index = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a heading
            if para.style and para.style.name.startswith('Heading'):
                heading_level = self._get_heading_level_from_style(para.style.name)
                sections.append(SectionSchema(
                    title=text,
                    level=heading_level,
                    order_index=len(sections)
                ))
            else:
                paragraphs.append(ParagraphSchema(
                    paragraph_id=str(uuid.uuid4()),
                    text=text,
                    order_index=paragraph_index
                ))
                paragraph_index += 1
        
        return paragraphs, sections
    
    def _extract_tables(self, doc: Document) -> List[TableSchema]:
        """Extract tables from DOCX.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of TableSchema objects
        """
        tables = []
        
        for table_idx, table in enumerate(doc.tables):
            if not table.rows or len(table.rows) < 2:
                continue
            
            # Extract headers (first row)
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            # Extract rows
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])
            
            # Generate CSV representation
            csv_lines = [','.join(headers)]
            csv_lines.extend([','.join(row) for row in rows])
            csv_representation = '\n'.join(csv_lines)
            
            tables.append(TableSchema(
                table_id=str(uuid.uuid4()),
                headers=headers,
                rows=rows,
                csv_representation=csv_representation,
                row_count=len(rows),
                column_count=len(headers),
                order_index=len(tables)
            ))
        
        return tables
    
    def _extract_raw_text(self, doc: Document) -> str:
        """Extract raw text from DOCX.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Raw text content
        """
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    
    def _normalize_text(self, paragraphs: List[ParagraphSchema]) -> str:
        """Normalize text from paragraphs.
        
        Args:
            paragraphs: List of paragraph schemas
            
        Returns:
            Normalized text
        """
        return '\n\n'.join([p.text for p in paragraphs])
    
    def _calculate_statistics(
        self,
        paragraphs: List[ParagraphSchema],
        tables: List[TableSchema],
        sections: List[SectionSchema]
    ) -> DocumentStatistics:
        """Calculate document statistics.
        
        Returns:
            DocumentStatistics object
        """
        word_count = sum(len(p.text.split()) for p in paragraphs)
        character_count = sum(len(p.text) for p in paragraphs)
        
        return DocumentStatistics(
            pages=1,  # DOCX is continuous, not paginated
            words=word_count,
            characters=character_count,
            paragraphs=len(paragraphs),
            tables=len(tables),
            images=0,
            sections=len(sections)
        )
    
    def _get_heading_level_from_style(self, style_name: str) -> int:
        """Extract heading level from style name.
        
        Args:
            style_name: Style name (e.g., 'Heading 1')
            
        Returns:
            Heading level (1-6)
        """
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 1
